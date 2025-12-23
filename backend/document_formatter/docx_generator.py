from __future__ import annotations

import logging
import re
import httpx
from backend.mermaid_renderer import render_mermaid_to_bytes
from io import BytesIO
from pathlib import Path
from typing import Dict, List, Any, Optional
from datetime import datetime

from backend.models import RequirementsResult, ExtractionResult

logger = logging.getLogger(__name__)


def _sanitize_mermaid_labels(diagram: str) -> str:
    """Sanitize Mermaid node labels to avoid parse errors in Kroki and mmdc.

    This wraps label text that contains special characters (commas, parens,
    slashes, colons) in double quotes when the label is inside square brackets
    and not already quoted. Example:
      A[Enterprise Systems (ERP, CRM)] -> A["Enterprise Systems (ERP, CRM)"]

    The function intentionally keeps existing quoted labels untouched.
    """
    if not diagram:
        return diagram

    def _quote_label(match):
        node = match.group(1)
        label = match.group(2)
        # If label already quoted, return as-is
        if (label.startswith('"') and label.endswith('"')) or (label.startswith("'") and label.endswith("'")):
            return match.group(0)
        # If label contains characters likely to break parsers, quote it
        if re.search(r'[(),:/\\\\]', label):
            # Escape any existing double quotes inside label
            safe = label.replace('"', '\\"')
            return f"{node}[\"{safe}\"]"
        return match.group(0)

    # Replace patterns like A[...label...] but avoid nested brackets
    result = re.sub(r"(\b[0-9A-Za-z_.$-]+)\[([^\]]+)\]", _quote_label, diagram)
    return result

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION_START
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. DOCX export will not work.")

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    logger.warning("PIL/Pillow not available. Logo may not display correctly.")


def clear_paragraph(paragraph):
    """Clear all runs from a paragraph by removing XML children."""
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def setup_styles(doc):
    """Configure document styles for clean, professional output."""
    # Normal / Body text style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    
    # Heading styles
    for level in (1, 2, 3, 4):
        try:
            h = doc.styles[f"Heading {level}"]
            h.font.name = "Calibri"
            h.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
            h.paragraph_format.space_after = Pt(6)
            h.paragraph_format.keep_with_next = True  # Avoids headings dangling at bottom
        except KeyError:
            logger.warning(f"Heading {level} style not found, skipping")
    
    try:
        list_bullet = doc.styles["List Bullet"]
        list_bullet.font.name = "Calibri"
        list_bullet.font.size = Pt(11)
        pf = list_bullet.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.first_line_indent = Inches(-0.25)
    except KeyError:
        logger.warning("List Bullet style not found, will use default")


def setup_page_formatting(doc, start_page_number: int = 1):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Set starting page number (for sections after front page and TOC)
    if start_page_number > 1:
        sect_pr = section._sectPr
        pg_num_type = OxmlElement('w:pgNumType')
        pg_num_type.set(qn('w:start'), str(start_page_number))
        sect_pr.append(pg_num_type)
    
    # Add page number to footer
    footer = section.footer
    # Ensure footer has at least one paragraph
    if len(footer.paragraphs) == 0:
        p = footer.add_paragraph()
    else:
        p = footer.paragraphs[0]
        # Clear existing content if any
        p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    run._r.append(fld)


def add_modern_front_page(doc, title: str, project_root: Optional[Path] = None):
    """Create a modern, professional front page with logo and title."""
    # Find logo path - try multiple locations (same path as used in frontend Header.jsx)
    if project_root is None:
        project_root = Path(__file__).parent.parent.parent
    
    logo_path = None
    # Try multiple possible logo locations (prioritize the same path as frontend)
    possible_logo_paths = [
        # Same path as frontend/src/components/Header.jsx uses
        project_root / "frontend" / "src" / "assets" / "logo-transparent.png",
        project_root / "frontend" / "src" / "assets" / "logo.png",
        # Docker/backend assets folder (copied during build)
        project_root / "assets" / "logo-transparent.png",
        project_root / "assets" / "logo.png",
        # Fallback locations
        project_root / "backend" / "assets" / "logo-transparent.png",
        project_root / "backend" / "assets" / "logo.png",
    ]
    
    for path in possible_logo_paths:
        if path.exists():
            logo_path = path
            logger.info(f"Found logo at: {logo_path}")
            break
    
    if logo_path is None:
        logger.warning(f"Logo not found. Tried paths: {[str(p) for p in possible_logo_paths]}")
    
    # Add spacing at top
    for _ in range(2):
        doc.add_paragraph()
    
    # Add logo if available
    if logo_path and logo_path.exists():
        try:
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add logo with appropriate size
            run = logo_para.add_run()
            if PIL_AVAILABLE:
                # Get image dimensions to maintain aspect ratio
                try:
                    img = Image.open(logo_path)
                    width, height = img.size
                    aspect_ratio = width / height
                    # Set height to 1.5 inches, calculate width
                    logo_height = Inches(1.5)
                    logo_width = Inches(1.5 * aspect_ratio)
                    logger.info(f"Logo dimensions: {width}x{height}, display size: {logo_width} x {logo_height}")
                except Exception as e:
                    logger.warning(f"Failed to get logo dimensions: {e}")
                    logo_width = Inches(3.5)
                    logo_height = Inches(1.5)
            else:
                logo_width = Inches(3.5)
                logo_height = Inches(1.5)
            
            run.add_picture(str(logo_path), width=logo_width, height=logo_height)
            logger.info(f"Successfully added logo to front page")
            
            # Add spacing after logo
            doc.add_paragraph()
            doc.add_paragraph()
        except Exception as e:
            logger.error(f"Failed to add logo to front page: {e}", exc_info=True)
    else:
        logger.warning("Logo not found or not accessible, skipping logo on front page")
    
    # Add title with modern styling (handle long titles with word wrap)
    # Adjust font size based on title length
    title_length = len(title)
    if title_length > 80:
        font_size = Pt(24)
    elif title_length > 60:
        font_size = Pt(28)
    else:
        font_size = Pt(32)
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(12)
    title_para.paragraph_format.space_before = Pt(0)
    
    # Enable word wrap and prevent truncation
    title_para.paragraph_format.widow_control = False
    title_para.paragraph_format.keep_together = False
    title_para.paragraph_format.keep_with_next = False
    
    # Set paragraph properties to allow wrapping and prevent truncation
    p_pr = title_para._element.get_or_add_pPr()
    # Remove any text overflow restrictions
    try:
        # Ensure text can wrap - set word wrap property
        wrap = OxmlElement('w:wordWrap')
        wrap.set(qn('w:val'), '0')  # 0 = wrap text
        p_pr.append(wrap)
        
        # Remove any overflow clip settings
        overflow = OxmlElement('w:overflowPunct')
        overflow.set(qn('w:val'), '0')  # Allow punctuation to overflow
        p_pr.append(overflow)
        
        # Remove any width restrictions - ensure paragraph can use full page width
        # Remove any indentation that might restrict width
        if p_pr.find(qn('w:ind')) is not None:
            p_pr.remove(p_pr.find(qn('w:ind')))
    except Exception:
        pass
    
    # Always add title as a single run to let Word handle wrapping naturally
    # This prevents truncation issues
    title_run = title_para.add_run(title)
    title_run.font.name = "Calibri"
    title_run.font.size = font_size
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(26, 84, 144)
    
    # Ensure the run doesn't have any restrictions that might cause truncation
    try:
        r_pr = title_run._element.get_or_add_rPr()
        # Remove any text effects that might truncate
        # Ensure no character limits are applied
    except Exception:
        pass
    
    logger.info(f"Added title to front page: '{title[:50]}...' (length: {title_length}, font size: {font_size})")
    
    # Add spacing
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Add date with subtle styling
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
    date_run.font.name = "Calibri"
    date_run.font.size = Pt(14)
    date_run.font.color.rgb = RGBColor(100, 100, 100)  # Gray
    
    # Add more spacing
    for _ in range(4):
        doc.add_paragraph()
    
    # Add company info section with modern styling
    company_para = doc.add_paragraph()
    company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_run = company_para.add_run("fusionAIx")
    company_run.font.name = "Calibri"
    company_run.font.size = Pt(20)
    company_run.font.bold = True
    company_run.font.color.rgb = RGBColor(26, 84, 144)
    
    # Add website
    website_para = doc.add_paragraph()
    website_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    website_run = website_para.add_run("www.fusionaix.com")
    website_run.font.name = "Calibri"
    website_run.font.size = Pt(12)
    website_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Add decorative line (optional - using spacing instead for cleaner look)
    doc.add_paragraph()


def add_manual_toc(doc, toc_entries: List[Dict[str, Any]]):
    """Add a manual table of contents without page numbers."""
    if not toc_entries:
        para = doc.add_paragraph("No table of contents entries available.")
        return
    
    for entry in toc_entries:
        para = doc.add_paragraph()
        para.style = 'Normal'
        
        # Add indentation based on level
        pf = para.paragraph_format
        indent_level = (entry.get('level', 1) - 1) * 0.5  # 0.5 inches per level
        if indent_level > 0:
            pf.left_indent = Inches(indent_level)
        
        # Add text only (no page numbers)
        text = entry.get('text', '')
        
        # Add text
        run = para.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def set_table_header_cell(cell):
    """Format a table header cell with bold text and styling."""
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            try:
                r.font.color.rgb = RGBColor(255, 255, 255)  # White text
            except:
                pass
    try:
        tc_pr = cell._element.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1a5490')
        shading.set(qn('w:val'), 'clear')
        tc_pr.append(shading)
    except Exception as e:
        logger.warning("Failed to set table header color: %s", e)


def finalize_table(table):
    """Apply final polish to a table: autofit, alignment, padding."""
    try:
        table.autofit = True
    except:
        pass
    
    # Set vertical alignment and cell padding
    for row in table.rows:
        for cell in row.cells:
            try:
                # Set vertical alignment to top for body cells
                tc_pr = cell._element.get_or_add_tcPr()
                v_align = OxmlElement('w:vAlign')
                v_align.set(qn('w:val'), 'top')
                tc_pr.append(v_align)
                
                # Set cell padding (top, right, bottom, left in twips - 1/20th of a point)
                tc_mar = OxmlElement('w:tcMar')
                for margin_name, margin_val in [('top', '80'), ('right', '80'), ('bottom', '80'), ('left', '80')]:
                    margin = OxmlElement(f'w:{margin_name}')
                    margin.set(qn('w:w'), margin_val)
                    margin.set(qn('w:type'), 'dxa')
                    tc_mar.append(margin)
                tc_pr.append(tc_mar)
            except Exception as e:
                logger.debug("Failed to set table cell formatting: %s", e)


def _add_formatted_text_to_paragraph(para, text: str):
    if not text:
        return
    
    if para.runs:
        clear_paragraph(para)
    
    parts = []
    last_end = 0
    
    for match in re.finditer(r'\*\*(.*?)\*\*', text):
        if match.start() > last_end:
            parts.append(('normal', text[last_end:match.start()]))
        parts.append(('bold', match.group(1)))
        last_end = match.end()
    
    if last_end < len(text):
        parts.append(('normal', text[last_end:]))
    
    if not parts:
        parts = [('normal', text)]
    
    for fmt_type, content in parts:
        if content:  # Only add non-empty content
            run = para.add_run(content)
            if fmt_type == 'bold':
                run.bold = True


def _add_bullet_paragraph(doc, content: str):
    """Add a paragraph with proper bullet point formatting using standard bullets."""
    para = doc.add_paragraph()
    
    # Set up proper indentation for bullet point
    pf = para.paragraph_format
    pf.left_indent = Inches(0.25)
    pf.first_line_indent = Inches(-0.25)
    pf.space_after = Pt(6)
    
    # Add bullet character
    bullet_run = para.add_run('• ')  # Standard bullet character
    bullet_run.font.name = "Calibri"
    bullet_run.font.size = Pt(11)
    
    # Add formatted content (handle bold formatting)
    parts = []
    last_end = 0
    
    for match in re.finditer(r'\*\*(.*?)\*\*', content):
        if match.start() > last_end:
            parts.append(('normal', content[last_end:match.start()]))
        parts.append(('bold', match.group(1)))
        last_end = match.end()
    
    if last_end < len(content):
        parts.append(('normal', content[last_end:]))
    
    if not parts:
        parts = [('normal', content)]
    
    for fmt_type, part_content in parts:
        if part_content:  # Only add non-empty content
            run = para.add_run(part_content)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            if fmt_type == 'bold':
                run.bold = True
    
    return para


def _start_table(doc, header_cells: List[str]):
    """Create a new table with header row. Returns the table object."""
    num_cols = len(header_cells)
    current_table = doc.add_table(rows=1, cols=num_cols)
    try:
        current_table.style = 'Light Grid Accent 1'
    except:
        try:
            current_table.style = 'Grid Table 1 Light'
        except:
            pass
    
    header_row_cells = current_table.rows[0].cells
    for col_idx, cell_text in enumerate(header_cells):
        cell = header_row_cells[col_idx]
        cell.text = cell_text
        set_table_header_cell(cell)
    
    return current_table


def _extract_headings_from_markdown(text: str) -> List[Dict[str, Any]]:
    headings = []
    lines = text.split('\n')
    
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('#### '):
            headings.append({'text': _clean_markdown_text(stripped[5:]), 'level': 4})
        elif stripped.startswith('### '):
            headings.append({'text': _clean_markdown_text(stripped[4:]), 'level': 3})
        elif stripped.startswith('## '):
            headings.append({'text': _clean_markdown_text(stripped[3:]), 'level': 2})
        elif stripped.startswith('# '):
            headings.append({'text': _clean_markdown_text(stripped[2:]), 'level': 1})
    
    return headings


def _parse_markdown_to_docx(doc, text: str):
    if not text:
        return
    
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        if re.match(r'^---+$', stripped):
            continue
        cleaned_lines.append(line)
    
    lines = cleaned_lines
    i = 0
    in_table = False
    current_table = None
    in_code_block = False
    code_block_lines = []
    code_block_lang = None
    last_was_blank = False
    list_item_buffer = []
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        # Detect unfenced Mermaid blocks: lines that begin with Mermaid diagram keywords
        mermaid_start_pattern = re.compile(r'^(?:flowchart|graph|sequenceDiagram|classDiagram|gantt|stateDiagram|pie|erDiagram)\b', re.IGNORECASE)
        if mermaid_start_pattern.match(stripped):
            # Collect contiguous mermaid lines until a blank line or a Caption: line
            block_lines = [line.rstrip('\n')]
            j = i + 1
            caption_text = None
            while j < len(lines):
                next_line = lines[j]
                next_stripped = next_line.strip()
                # Stop if blank line; allow Caption: line to be captured separately
                if not next_stripped:
                    break
                if re.match(r'^Caption:\s*', next_stripped, flags=re.IGNORECASE):
                    caption_text = re.sub(r'^Caption:\s*', '', next_stripped, flags=re.IGNORECASE).strip()
                    j += 1
                    break
                # If another markdown delimiter or heading starts, stop
                if next_stripped.startswith('```') or next_stripped.startswith('#'):
                    break
                block_lines.append(next_line.rstrip('\n'))
                j += 1

            block_text = '\n'.join(block_lines)
            # Sanitize labels to avoid Kroki/mmdc parse errors (unquoted commas/paren etc.)
            sanitized_block = _sanitize_mermaid_labels(block_text)

            # Try to render same as fenced block handling
            rendered = None
            try:
                rendered = render_mermaid_to_bytes(sanitized_block, fmt='png')
            except FileNotFoundError:
                logger.info('Local mmdc not found; will fall back to Kroki')
            except Exception as e:
                logger.exception('Local mmdc rendering failed: %s', e)

            if rendered:
                try:
                    logger.info('Inserting locally rendered mermaid image into DOCX (unfenced)')
                    img_stream = BytesIO(rendered)
                    doc.add_picture(img_stream, width=Inches(5))
                    if caption_text:
                        cap_para = doc.add_paragraph()
                        cap_run = cap_para.add_run(caption_text)
                        cap_run.italic = True
                        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    logger.warning('Failed to insert locally rendered mermaid image into docx (unfenced): %s', e)
                    rendered = None

            if not rendered:
                # Kroki fallback
                try:
                    kroki_url = 'https://kroki.io/mermaid/png'
                    resp = httpx.post(kroki_url, content=sanitized_block.encode('utf-8'), headers={"Content-Type": "text/plain"}, timeout=30.0)
                    if resp.status_code == 200:
                        img_bytes = resp.content
                        try:
                            img_stream = BytesIO(img_bytes)
                            doc.add_picture(img_stream, width=Inches(5))
                            if caption_text:
                                cap_para = doc.add_paragraph()
                                cap_run = cap_para.add_run(caption_text)
                                cap_run.italic = True
                                cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception as e:
                            logger.warning('Failed to insert Kroki mermaid image into docx (unfenced): %s', e)
                            para = doc.add_paragraph(style='Normal')
                            para.style.font.name = 'Consolas'
                            para.style.font.size = Pt(10)
                            para.text = block_text
                    else:
                        logger.warning('Kroki returned status %s for mermaid render (unfenced)', resp.status_code)
                        para = doc.add_paragraph(style='Normal')
                        para.style.font.name = 'Consolas'
                        para.style.font.size = Pt(10)
                        para.text = block_text
                except Exception as e:
                    logger.exception('Failed to call Kroki for mermaid rendering (unfenced): %s', e)
                    para = doc.add_paragraph(style='Normal')
                    para.style.font.name = 'Consolas'
                    para.style.font.size = Pt(10)
                    para.text = block_text

            i = j
            continue
        
        # Handle code blocks
        if stripped.startswith('```'):
            # Determine language if provided: ```lang
            lang = stripped[3:].strip().lower()
            if in_code_block:
                # End code block
                if code_block_lines:
                    block_text = '\n'.join(code_block_lines)
                    # Detect and strip an inline trailing caption if present (common LLM output quirk)
                    caption_text = None
                    m_caption = re.search(r"\n?\s*Caption:\s*(.+)\s*$", block_text, flags=re.IGNORECASE)
                    if m_caption:
                        caption_text = m_caption.group(1).strip()
                        block_text = re.sub(r"\n?\s*Caption:\s*.+\s*$", "", block_text, flags=re.IGNORECASE)

                    if code_block_lang == 'mermaid':
                        # Prefer to render locally via mmdc (Mermaid CLI). If unavailable
                        # or rendering fails, fall back to Kroki. If both fail, insert raw code.
                        rendered = None
                        # Sanitize fenced mermaid block before rendering
                        sanitized_block = _sanitize_mermaid_labels(block_text)
                        try:
                            rendered = render_mermaid_to_bytes(sanitized_block, fmt='png')
                        except FileNotFoundError:
                            logger.info('Local mmdc not found; will fall back to Kroki')
                        except Exception as e:
                            logger.exception('Local mmdc rendering failed: %s', e)

                        if rendered:
                            try:
                                logger.info('Inserting locally rendered mermaid image into DOCX')
                                img_stream = BytesIO(rendered)
                                doc.add_picture(img_stream, width=Inches(5))
                                # Insert caption if present
                                if caption_text:
                                    cap_para = doc.add_paragraph()
                                    cap_run = cap_para.add_run(caption_text)
                                    cap_run.italic = True
                                    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            except Exception as e:
                                logger.warning('Failed to insert locally rendered mermaid image into docx: %s', e)
                                rendered = None

                        if not rendered:
                            # Try Kroki as a fallback
                            try:
                                kroki_url = 'https://kroki.io/mermaid/png'
                                resp = httpx.post(kroki_url, content=sanitized_block.encode('utf-8'), headers={"Content-Type": "text/plain"}, timeout=30.0)
                                if resp.status_code == 200:
                                    logger.info('Kroki returned image bytes for mermaid block')
                                    img_bytes = resp.content
                                    try:
                                        img_stream = BytesIO(img_bytes)
                                        doc.add_picture(img_stream, width=Inches(5))
                                        if caption_text:
                                            cap_para = doc.add_paragraph()
                                            cap_run = cap_para.add_run(caption_text)
                                            cap_run.italic = True
                                            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    except Exception as e:
                                        logger.warning('Failed to insert Kroki mermaid image into docx: %s', e)
                                        # Fall through to insert raw code
                                        para = doc.add_paragraph(style='Normal')
                                        para.style.font.name = 'Consolas'
                                        para.style.font.size = Pt(10)
                                        para.text = block_text
                                else:
                                    logger.warning('Kroki returned status %s for mermaid render', resp.status_code)
                                    para = doc.add_paragraph(style='Normal')
                                    para.style.font.name = 'Consolas'
                                    para.style.font.size = Pt(10)
                                    para.text = block_text
                            except Exception as e:
                                logger.exception('Failed to call Kroki for mermaid rendering: %s', e)
                                para = doc.add_paragraph(style='Normal')
                                para.style.font.name = 'Consolas'
                                para.style.font.size = Pt(10)
                                para.text = block_text
                    else:
                        para = doc.add_paragraph(style='Normal')
                        para.style.font.name = 'Consolas'
                        para.style.font.size = Pt(10)
                        para.text = block_text
                code_block_lines = []
                in_code_block = False
                code_block_lang = None
            else:
                # Start code block
                in_code_block = True
                code_block_lang = lang or None
            i += 1
            continue
        
        table_header_pattern = re.compile(r"^\s*\|.*\|\s*$")
        table_sep_pattern = re.compile(r"^\s*\|?\s*[:\-]+(?:\s*\|\s*[:\-]+)+\s*\|?\s*$")
        if i + 1 < len(lines) and table_header_pattern.match(line) and table_sep_pattern.match(lines[i + 1].strip()):
            header_line = line.strip().strip('|')
            header_cells = [h.strip() for h in re.split(r'\s*\|\s*', header_line)]
            try:
                current_table = _start_table(doc, header_cells)
                in_table = True
            except Exception as e:
                logger.warning("Failed to start table: %s", e)
                in_table = False
                current_table = None

            j = i + 2
            while j < len(lines):
                row_line = lines[j].strip()
                if not row_line or not row_line.startswith('|'):
                    break
                row_cells = [c.strip() for c in re.split(r'\s*\|\s*', row_line.strip().strip('|'))]
                try:
                    row = current_table.add_row()
                    for col_idx, cell_text in enumerate(row_cells[: len(header_cells) ]):
                        try:
                            cell = row.cells[col_idx]
                            cell.text = _clean_markdown_text(cell_text)
                        except Exception:
                            pass
                except Exception as e:
                    logger.debug("Failed to add table row: %s", e)
                j += 1

            if in_table and current_table:
                finalize_table(current_table)
            in_table = False
            current_table = None
            i = j
            continue

        if in_table and current_table:
            finalize_table(current_table)
            in_table = False
            current_table = None
        
        if stripped.startswith('- ') or stripped.startswith('* ') or re.match(r'^\d+\.\s', stripped):
            if stripped.startswith('- ') or stripped.startswith('* '):
                content = stripped[2:].strip()
                is_bullet = True
            else:
                content = re.sub(r'^\d+\.\s', '', stripped).strip()
                is_bullet = False
            
            if i + 1 < len(lines):
                next_line = lines[i + 1]
                next_stripped = next_line.strip()
                if (next_stripped and 
                    not next_stripped.startswith('- ') and 
                    not next_stripped.startswith('* ') and 
                    not re.match(r'^\d+\.\s', next_stripped)):
                    list_item_buffer.append(content)
                    list_item_buffer.append(next_stripped)
                    i += 2
                    while i < len(lines):
                        cont_line = lines[i]
                        cont_stripped = cont_line.strip()
                        if (not cont_stripped or 
                            cont_stripped.startswith('- ') or 
                            cont_stripped.startswith('* ') or 
                            re.match(r'^\d+\.\s', cont_stripped)):
                            break
                        list_item_buffer.append(cont_stripped)
                        i += 1
                    # Now add the complete list item
                    full_content = ' '.join(list_item_buffer)
                    list_item_buffer = []
                    full_content = _capitalize_sentence(full_content)
                    if is_bullet:
                        _add_bullet_paragraph(doc, full_content)
                    else:
                        para = doc.add_paragraph(style='List Number')
                        _add_formatted_text_to_paragraph(para, full_content)
                    continue
            
            # Single-line list item
            content = _capitalize_sentence(content)
            if is_bullet:
                _add_bullet_paragraph(doc, content)
            else:
                para = doc.add_paragraph(style='List Number')
                _add_formatted_text_to_paragraph(para, content)
            i += 1
            continue
        
        # Regular text/headings
        _add_text_line(doc, stripped)
        i += 1
    
    # Finalize any remaining table
    if in_table and current_table:
        finalize_table(current_table)
    
    # Handle any remaining code block
    if in_code_block and code_block_lines:
        para = doc.add_paragraph(style='Normal')
        para.style.font.name = 'Consolas'
        para.style.font.size = Pt(10)
        para.text = '\n'.join(code_block_lines)


def _clean_markdown_text(text: str) -> str:
    if not text:
        return text

    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*\*([^*]*)\*\*', r'\1', text)
    text = re.sub(r'(?<!\*)\*([^*]+)\*(?!\*)', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'^---+$', '', text, flags=re.MULTILINE)
    text = re.sub(r'[ \t]+', ' ', text)
    text = text.strip()
    return text


def _capitalize_sentence(text: str) -> str:
    if not text:
        return text
    text = text.strip()
    if text and text[0].islower():
        return text[0].upper() + text[1:]
    return text


def _add_text_line(doc, line: str):
    """Add a single line of text to the document, handling headings and regular text."""
    stripped = line.strip()
    if not stripped:
        return
    
    if re.match(r'^---+$', stripped):
        return
    
    if stripped.startswith('#### '):
        header_text = _clean_markdown_text(stripped[5:])
        if header_text:
            doc.add_heading(header_text, 4)
    elif stripped.startswith('### '):
        header_text = _clean_markdown_text(stripped[4:])
        if header_text:
            doc.add_heading(header_text, 3)
    elif stripped.startswith('## '):
        header_text = _clean_markdown_text(stripped[3:])
        if header_text:
            doc.add_heading(header_text, 2)
    elif stripped.startswith('# '):
        header_text = _clean_markdown_text(stripped[2:])
        if header_text:
            doc.add_heading(header_text, 1)
    else:
        content = _clean_markdown_text(stripped)
        if content:
            if content and len(content) > 0 and content[0].islower() and content[0].isalpha():
                content = _capitalize_sentence(content)
            para = doc.add_paragraph()
            _add_formatted_text_to_paragraph(para, content)


def generate_rfp_docx(
    individual_responses: List[Dict[str, Any]],
    requirements_result: RequirementsResult,
    extraction_result: ExtractionResult,
    rfp_title: Optional[str] = None,
    output_path: Optional[Path] = None,
) -> bytes:
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is not installed. Install it with: pip install python-docx")
    
    logger.info("Generating DOCX document: %d requirement responses", len(individual_responses))
    
    # Extract TOC entries before generating content
    is_structured = len(individual_responses) == 1 and individual_responses[0].get('requirement_id') == 'STRUCTURED'
    toc_entries = []
    
    if is_structured:
        # Prefer detected sections from structure detection if available
        if (requirements_result.structure_detection and 
            requirements_result.structure_detection.detected_sections):
            # Use detected sections from structure detection
            # Format with numbers (e.g., "1. Executive Summary", "2. Understanding...")
            toc_entries = [
                {'text': f"{idx + 1}. {section}", 'level': 1}
                for idx, section in enumerate(requirements_result.structure_detection.detected_sections)
            ]
        else:
            # Fallback: Extract headings from structured response
            response_text = individual_responses[0].get('response', '')
            headings = _extract_headings_from_markdown(response_text)
            # Filter to only include numbered sections (e.g., "1. Executive Summary", "2. Understanding...")
            # Exclude non-numbered headings like "Ministry of Trade and Industry (MTI)"
            numbered_pattern = re.compile(r'^\d+\.\s+')
            toc_entries = [
                h for h in headings 
                if h['level'] <= 2 and numbered_pattern.match(h['text'])
            ]
    else:
        # Create TOC entries for each requirement
        for idx, resp_data in enumerate(individual_responses, 1):
            req_id = resp_data.get('requirement_id', 'N/A')
            toc_entries.append({
                'text': f"Requirement {idx}: {req_id}",
                'level': 2,
            })
    
    doc = Document()
    
    # Setup styles first (cleanest output)
    setup_styles(doc)
    
    # Setup page formatting for first section (front page - no page numbers)
    setup_page_formatting(doc, start_page_number=1)
    
    # Disable page numbers for first section (front page)
    section = doc.sections[0]
    sect_pr = section._sectPr
    # Remove footer for first section
    footer = section.footer
    # Clear footer paragraphs instead of using clear_content()
    for para in footer.paragraphs[:]:
        p = para._element
        p.getparent().remove(p)
    
    # Determine project root for logo path (always calculate from file location)
    project_root = Path(__file__).parent.parent.parent
    
    # Create modern front page
    final_title = rfp_title or extraction_result.language.upper()
    add_modern_front_page(doc, final_title, project_root)
    
    # Add new section for TOC (no page numbers) - NEW_PAGE creates the page break
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    section.is_linked_to_previous = False  # Break link to previous section's header/footer
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    sect_pr = section._sectPr
    footer = section.footer
    # Clear footer paragraphs instead of using clear_content()
    for para in footer.paragraphs[:]:
        p = para._element
        p.getparent().remove(p)
    
    doc.add_heading("Table of Contents", 1)
    doc.add_paragraph()  # Add spacing after heading
    
    # Add manual TOC with structure or requirement names (no page numbers)
    add_manual_toc(doc, toc_entries)
    
    # Add new section for main content (with page numbers starting from 1, but it's actually page 3) - NEW_PAGE creates the page break
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    section.is_linked_to_previous = False  # Break link to previous section's header/footer
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    
    # Set page number to start at 1 for this section (even though it's the 3rd physical page)
    sect_pr = section._sectPr
    # Remove any existing pgNumType element
    existing_pg_num = sect_pr.find(qn('w:pgNumType'))
    if existing_pg_num is not None:
        sect_pr.remove(existing_pg_num)
    # Create new page number type with restart at 1
    pg_num_type = OxmlElement('w:pgNumType')
    pg_num_type.set(qn('w:start'), '1')  # Start numbering at 1
    pg_num_type.set(qn('w:fmt'), 'decimal')  # Use decimal numbering (1, 2, 3...)
    sect_pr.append(pg_num_type)
    
    # Setup page formatting with footer for this section
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Add page number to footer
    footer = section.footer
    # Ensure footer has at least one paragraph
    if len(footer.paragraphs) == 0:
        p = footer.add_paragraph()
    else:
        p = footer.paragraphs[0]
        # Clear existing content if any
        p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    run._r.append(fld)
    
    doc.add_heading("Company Overview", 1)
    overview_text = """At fusionAIx, we believe that the future of digital transformation lies in the seamless blend of low-code platforms and artificial intelligence. Our core team brings together decades of implementation experience, domain expertise, and a passion for innovation. We partner with enterprises to reimagine processes, accelerate application delivery, and unlock new levels of efficiency. We help businesses scale smarter, faster, and with greater impact.

With a collaborative spirit and a commitment to excellence, our team transforms complex challenges into intelligent, practical solutions. fusionAIx is not just about technology—it's about empowering people, industries, and enterprises to thrive in a digital-first world.

We are proud to be officially recognized as a Great Place To Work® Certified Company for 2025–26, reflecting our commitment to a culture built on trust, innovation, and people-first values.

fusionAIx delivers tailored solutions that blend AI and automation to drive measurable results across industries. We are a niche Pega partner with 20+ successful Pega Constellation implementations across the globe. As Constellation migration experts, we focus on pattern-based development with Constellation, enabling faster project go-lives than traditional implementation approaches.

Our proven capabilities span three core technology platforms: Pega Constellation, Microsoft Power Platform, and ServiceNow. Through these platforms, we provide comprehensive services including Low Code/No Code development, Digital Process Transformation, and AI & Data solutions.

To accelerate time-to-value, fusionAIx offers proprietary accelerators and solution components including fxAgentSDK, fxAIStudio, fxMockUpToView, and fxSmartDCO. These tools enable rapid development, intelligent automation, and streamlined project delivery.

We support clients across diverse industries including Insurance, Banking & Finance, Government & Public Sector, Automotive & Fleet Management, and Travel & Tourism, combining platform expertise with structured knowledge transfer to help customers build sustainable, future-ready capabilities."""
    
    for para_text in overview_text.split('\n\n'):
        if para_text.strip():
            doc.add_paragraph(para_text.strip())
    
    doc.add_page_break()
    
    is_structured = len(individual_responses) == 1 and individual_responses[0].get('requirement_id') == 'STRUCTURED'
    
    if is_structured:
        response_text = individual_responses[0].get('response', '')
        _parse_markdown_to_docx(doc, response_text)
    else:
        doc.add_heading("Solution Requirement Responses", 1)
        for idx, resp_data in enumerate(individual_responses, 1):
            req_heading = doc.add_heading(f"Requirement {idx}: {resp_data.get('requirement_id', 'N/A')}", 2)
            
            req_para = doc.add_paragraph()
            req_para.add_run("Requirement: ").bold = True
            req_text = resp_data.get('requirement_text', '')
            if req_text:
                req_para.add_run(_capitalize_sentence(req_text))
            
            resp_heading = doc.add_heading("Response", 3)
            _parse_markdown_to_docx(doc, resp_data.get('response', ''))
            
            if resp_data.get('quality'):
                quality = resp_data['quality']
                quality_para = doc.add_paragraph()
                quality_para.add_run(f"Quality Score: {quality.get('score', 0):.0f}/100 | ").bold = True
                quality_para.add_run(f"Completeness: {quality.get('completeness', 'unknown')} | ")
                quality_para.add_run(f"Relevance: {quality.get('relevance', 'unknown')}")
            
            doc.add_paragraph()
    
    if output_path:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        logger.info("DOCX saved to: %s", output_path.absolute())
        return output_path.read_bytes()
    else:
        # Use BytesIO instead of temp file round-trip
        buf = BytesIO()
        doc.save(buf)
        bytes_data = buf.getvalue()
        logger.info("DOCX generated: %d bytes (in memory)", len(bytes_data))
        return bytes_data

